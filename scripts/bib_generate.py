
from docx import Document

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = Document()

# Function to set the page margins to ensure numbers cover more space
def set_page_margins(section, top=1, right=1, bottom=1, left=1):
    section.top_margin = Inches(top)  # Set margins in inches
    section.right_margin = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)

# Function to add two numbers on a single page
def add_two_numbers_page(number1, number2):
    # Set up the first half of the page
    p1 = doc.add_paragraph()
    run1 = p1.add_run(f'{number1:04d}')
    run1.font.size = Pt(150)  # Large font size for the number
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the text
    
    # Add a paragraph break with specific spacing to center the first number in its half
    p1_format = p1.paragraph_format
    p1_format.space_after = Pt(200)  # Adjust this value as needed to center vertically

    # Set up the second half of the page
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'{number2:04d}')
    run2.font.size = Pt(150)  # Large font size for the number
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the text
    
    # Add a paragraph break with specific spacing to center the second number in its half
    p2_format = p2.paragraph_format
    p2_format.space_before = Pt(200)  # Adjust this value as needed to center vertically

    # Add a page break after each set of numbers except the last one
    if number2 < 1200:
        doc.add_page_break()

# Set margins for the first section
set_page_margins(doc.sections[0])

# Add numbers from 0001 to 1200, two per page
for number in range(1, 1201, 2):
    add_two_numbers_page(number, number + 1)

# Save the document
doc.save('marathon_numbers.docx')

print('Document created successfully: marathon_numbers.docx')
